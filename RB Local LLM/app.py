from langchain_community.llms import Ollama
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnablePassthrough
from langchain_core.output_parsers import StrOutputParser
import docx

# from langchain.output_parsers import StrOutputParser
from langchain_community.embeddings.fastembed import FastEmbedEmbeddings
from langchain_community.vectorstores import Chroma

# from langchain.chains import RetrievalQA
import streamlit as st
import os
from PyPDF2 import PdfReader
from langchain_core.documents import Document
from bs4 import BeautifulSoup
import requests

# Initialize tqdm lock
from tqdm import tqdm

tqdm(disable=True, total=0)
tqdm.write("test")

# Set up current directory and file tracker
dir_path = os.path.dirname(os.path.realpath(__file__))
print(f"Current working directory: {dir_path}")
path_for_file_tracker = os.path.join(dir_path, "list_of_files.txt")
if not os.path.isfile(path_for_file_tracker):
    with open(path_for_file_tracker, "w") as file:
        file.write("List of files in the embeddings database:\n")
    print("File tracker created")

# Create directory for storing the ChromaDB vector store
directory_path = "chroma_langchain_db"
if not os.path.isdir(directory_path):
    os.mkdir(directory_path)

# Set up Streamlit page configuration
st.set_page_config(
    page_title="Chat with your documents",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# Initialize model, embeddings, and text splitter
cached_llm = Ollama(model="llama3.1")
embedding = FastEmbedEmbeddings()
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=1024, chunk_overlap=80, length_function=len
)

# Initialize the vector store
vector_store = Chroma(embedding_function=embedding, persist_directory=directory_path)


# Function to check if the vector store is empty
def is_vector_store_empty(vector_store: Chroma):
    try:
        results = vector_store.similarity_search("test", k=1)
        return len(results) == 0
    except Exception:
        return True


# Functions to process and add documents to the vector store
def add_chunks_text_docx(docx_file, current_vector_store: Chroma):

    doc = docx.Document(docx_file)
    text_content = "\n".join([para.text for para in doc.paragraphs])
    document = Document(page_content=text_content)
    chunks = text_splitter.split_documents([document])
    current_vector_store.add_documents(chunks)


def add_chunks_text_txt(text_file, current_vector_store: Chroma):
    text_content = text_file.read() if hasattr(text_file, "read") else str(text_file)
    document = Document(page_content=text_content)
    chunks = text_splitter.split_documents([document])
    current_vector_store.add_documents(chunks)


def add_chunks_text_pdf(pdf_docs, current_vector_store: Chroma):
    text = ""
    pdf_reader = PdfReader(pdf_docs)
    for page in pdf_reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n\n"
    document = Document(page_content=text)
    chunks = text_splitter.split_documents([document])
    current_vector_store.add_documents(chunks)


# Function to extract text from HTML (e.g., legislation websites)
def add_chunks_text_html(url, current_vector_store: Chroma):
    response = requests.get(url)
    if response.status_code == 200:
        soup = BeautifulSoup(response.content, "html.parser")
        text_content = soup.get_text(separator="\n", strip=True)
        document = Document(page_content=text_content)
        chunks = text_splitter.split_documents([document])
        current_vector_store.add_documents(chunks)
        print(f"Successfully processed and added content from {url}")
    else:
        print(f"Failed to retrieve {url}. Status code: {response.status_code}")


# Process user questions
def process_question(question: str, vector_db: Chroma) -> str:
    if is_vector_store_empty(vector_db):
        return st.warning(
            "The vector store is empty. Please upload a document or provide a URL first."
        )

    llm = Ollama(model="llama3.1", temperature=0.0)

    retriever = vector_db.as_retriever(
        search_type="similarity_score_threshold",
        search_kwargs={"k": 20, "score_threshold": 0.1},
    )

    template = """You are an AI language model assistant. Your Task is to generate a precise answer to the question based ONLY on the following context:
    {context}
    Question: {question}
    If you don't know the answer, just say you don't know. Do not attempt to provide an answer based on assumptions or external knowledge.
    Provide references or snippets from the context to support your response."""

    prompt = ChatPromptTemplate.from_template(template)

    chain = (
        {"context": retriever, "question": RunnablePassthrough()}
        | prompt
        | llm
        | StrOutputParser()
    )

    response = chain.invoke(question)
    return response


# Main function for Streamlit app
def main():
    st.subheader("Chat with your documents", divider="gray")

    # Dynamically join the path to your logo image
    # logo_path = os.path.join(os.getcwd(), "RB-Switchgear-Group-Logo.png")

    # # Check if the file exists
    # if os.path.exists(logo_path):
    #     st.image(logo_path, caption="Confidential", width=230)
    # else:
    #     st.warning(f"Logo image not found at: {logo_path}")

    col1, col2 = st.columns([1, 2])

    if "messages" not in st.session_state:
        st.session_state["messages"] = []

    if "vector_db" not in st.session_state:
        st.session_state["vector_db"] = vector_store

    # File uploader and URL input in the left column
    file_upload = col1.file_uploader(
        "Upload a PDF, DOCX, TXT file", type=["pdf", "docx", "txt", "doc"]
    )
    url_input = col1.text_input(
        "Or enter a URL to extract data from (e.g., legislation websites)"
    )

    # Process file uploads
    if file_upload is not None:
        with st.spinner(f"Uploading {file_upload.name}..."):
            if file_upload.type == "text/plain":
                add_chunks_text_txt(file_upload, st.session_state["vector_db"])
            elif file_upload.type == "application/pdf":
                add_chunks_text_pdf(file_upload, st.session_state["vector_db"])
            elif (
                file_upload.type
                == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ):
                add_chunks_text_docx(file_upload, st.session_state["vector_db"])
            with open(path_for_file_tracker, "a") as myfile:
                myfile.write("\n" + str(file_upload.name))
        file_upload = None

    # Process URL inputs
    if url_input:
        with st.spinner(f"Processing URL {url_input}..."):
            add_chunks_text_html(url_input, st.session_state["vector_db"])
            with open(path_for_file_tracker, "a") as myfile:
                myfile.write("\n" + str(url_input))

    # Display a warning if the vector store is empty
    elif file_upload is None and is_vector_store_empty(vector_store):
        st.warning("Upload a file or provide a URL to begin.")

    # Right column for chatbot interaction
    with col2:
        message_container = st.container()

        for message in st.session_state["messages"]:
            avatar = "🤖" if message["role"] == "assistant" else "😎"
            with message_container.chat_message(message["role"], avatar=avatar):
                st.markdown(message["content"])

        if prompt := st.chat_input("Enter a prompt here..."):
            try:
                st.session_state["messages"].append({"role": "user", "content": prompt})
                message_container.chat_message("user", avatar="😎").markdown(prompt)
                with message_container.chat_message("assistant", avatar="🤖"):
                    with st.spinner("Processing..."):
                        if not is_vector_store_empty(st.session_state["vector_db"]):
                            response = process_question(
                                prompt, st.session_state["vector_db"]
                            )
                            st.markdown(response)
                        else:
                            st.warning(
                                "Vector store is empty. Please upload a document or provide a URL."
                            )
                st.session_state["messages"].append(
                    {"role": "assistant", "content": response}
                )
            except Exception as e:
                st.error(f"Error: {e}")


    """------------------------------Testing Zone-------------------------"""
    from deepeval.dataset import EvaluationDataset
    from deepeval.metrics import ContextualRecallMetric


    metric = ContextualRecallMetric(threshold=0.5)

    dataset = EvaluationDataset(test_cases=[...])
    dataset.evaluate(metrics=[metric])


    from deepeval.synthesizer import Synthesizer

    synthesizer = Synthesizer()
    goldens = synthesizer.generate_goldens_from_docs(
        document_paths=[
            "pdfs/health-and-safety-risk-management.pdf",
            "pdfs/ECI-SHE7-SHE-Guide.pdf",
            "pdfs/hsg129.pdf",
            "pdfs/l153.pdf",
            "pdfs/the-business-case-for-engineering-in-health-and-safety.pdf",
        ],
        max_goldens_per_context=15,
        include_expected_output=True,
    )

    from deepeval.metrics import GEval, ContextualRecallMetric
    from deepeval.test_case import LLMTestCaseParams

    contextual_recall = ContextualRecallMetric(threshold=0.6)
    correctness = GEval(
        name="Correctness",
        criteria="Determine whether the actual output is factually correct based on the expected output.",
        evaluation_params=[
            LLMTestCaseParams.ACTUAL_OUTPUT,
            LLMTestCaseParams.EXPECTED_OUTPUT,
        ],
        threshold=0.6,
    )

    """Run the benchmark by first converting the goldens generated in Step 1 to test cases that are reday for evaluation"""
    from deepeval.test_case import LLMTestCase
    from deepeval.dataset import EvaluationDataset

    ...

    test_cases = []
    for golden in goldens:
        query = golden.input
        llm_output, retrieved_nodes = process_question(query, vector_db)
        test_case = LLMTestCase(
            input=query,
            actual_output=llm_output,
            retrieval_context=retrieved_nodes,
            expected_output=golden.expected_output,
        )

    test_cases.append(test_case)

    dataset = EvaluationDataset(test_cases=test_cases)
    dataset.evaluate(metrics=[correctness, contextual_recall])
    """------------------------------Testing Zone-------------------------"""
if __name__ == "__main__":
    main()
